import zipfile
from pathlib import Path
from typing import Any

import pymupdf
import pytest
from docx import Document


# ── PDF fixtures ─────────────────────────────────────────────


@pytest.fixture
def multi_page_pdf(tmp_path: Path) -> Path:
    """Create a 3-page PDF with known text content."""
    path = tmp_path / "multi_page.pdf"
    doc = pymupdf.open()
    for i in range(1, 4):
        page = doc.new_page()
        page.insert_text((72, 72), f"Content of page {i}.")
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def pdf_with_empty_page(tmp_path: Path) -> Path:
    """Create a 3-page PDF where page 2 has no text (simulates image-only page)."""
    path = tmp_path / "empty_page.pdf"
    doc = pymupdf.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Page one has text.")
    doc.new_page()  # page 2: blank
    page3 = doc.new_page()
    page3.insert_text((72, 72), "Page three has text.")
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def invalid_pdf(tmp_path: Path) -> Path:
    """Create a file with .pdf extension but invalid content."""
    path = tmp_path / "invalid.pdf"
    path.write_text("This is not a PDF file.")
    return path


# ── EPUB fixtures ────────────────────────────────────────────

_CONTAINER_XML = """\
<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>"""


def _make_opf(chapter_ids: list[str]) -> str:
    manifest_items = "\n    ".join(
        f'<item id="{cid}" href="{cid}.xhtml" media-type="application/xhtml+xml"/>'
        for cid in chapter_ids
    )
    spine_items = "\n    ".join(
        f'<itemref idref="{cid}"/>' for cid in chapter_ids
    )
    return f"""\
<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>Test Book</dc:title>
  </metadata>
  <manifest>
    {manifest_items}
  </manifest>
  <spine>
    {spine_items}
  </spine>
</package>"""


def _make_xhtml(body_content: str) -> str:
    return f"""\
<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>Chapter</title></head>
<body>
{body_content}
</body>
</html>"""


def _build_epub(
    path: Path,
    chapters: dict[str, str],
    *,
    include_encryption: bool = False,
) -> Path:
    chapter_ids = list(chapters.keys())
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip")
        zf.writestr("META-INF/container.xml", _CONTAINER_XML)
        if include_encryption:
            zf.writestr("META-INF/encryption.xml", "<encryption/>")
        zf.writestr("OEBPS/content.opf", _make_opf(chapter_ids))
        for cid, body in chapters.items():
            zf.writestr(f"OEBPS/{cid}.xhtml", _make_xhtml(body))
    return path


@pytest.fixture
def multi_chapter_epub(tmp_path: Path) -> Path:
    return _build_epub(
        tmp_path / "multi_chapter.epub",
        {
            "ch1": "<p>Chapter 1 content goes here.</p>",
            "ch2": "<p>Chapter 2 content goes here.</p>",
            "ch3": "<p>Chapter 3 content goes here.</p>",
        },
    )


@pytest.fixture
def single_chapter_epub(tmp_path: Path) -> Path:
    return _build_epub(
        tmp_path / "single_chapter.epub",
        {"ch1": "<p>Only chapter with some text.</p>"},
    )


@pytest.fixture
def epub_with_whitespace_chapter(tmp_path: Path) -> Path:
    return _build_epub(
        tmp_path / "whitespace_chapter.epub",
        {
            "ch1": "<p>Chapter 1 text.</p>",
            "ch2": "   \n   ",
            "ch3": "<p>Chapter 3 text.</p>",
        },
    )


@pytest.fixture
def drm_protected_epub(tmp_path: Path) -> Path:
    return _build_epub(
        tmp_path / "drm.epub",
        {"ch1": "<p>Protected content.</p>"},
        include_encryption=True,
    )


@pytest.fixture
def invalid_epub(tmp_path: Path) -> Path:
    path = tmp_path / "invalid.epub"
    path.write_text("This is not an EPUB file.")
    return path


@pytest.fixture
def epub_with_no_chapters(tmp_path: Path) -> Path:
    path = tmp_path / "no_chapters.epub"
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip")
        zf.writestr("META-INF/container.xml", _CONTAINER_XML)
        zf.writestr("OEBPS/content.opf", _make_opf([]))
    return path


# ── DOCX fixtures ────────────────────────────────────────────


def _add_heading(doc: Any, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_paragraph(doc: Any, text: str) -> None:
    doc.add_paragraph(text)


def _add_table(doc: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text


@pytest.fixture
def multi_section_docx(tmp_path: Path) -> Path:
    """DOCX with H1 and H2 headings creating multiple sections."""
    path = tmp_path / "multi_section.docx"
    doc = Document()
    _add_paragraph(doc, "Introduction text before any heading.")
    _add_heading(doc, "Chapter One", level=1)
    _add_paragraph(doc, "Chapter one content here.")
    _add_heading(doc, "Section 1.1", level=2)
    _add_paragraph(doc, "Section 1.1 content here.")
    _add_heading(doc, "Chapter Two", level=1)
    _add_paragraph(doc, "Chapter two content here.")
    doc.save(str(path))
    return path


@pytest.fixture
def docx_no_headings(tmp_path: Path) -> Path:
    """DOCX with only paragraphs, no headings."""
    path = tmp_path / "no_headings.docx"
    doc = Document()
    _add_paragraph(doc, "First paragraph of text.")
    _add_paragraph(doc, "Second paragraph of text.")
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_tables(tmp_path: Path) -> Path:
    """DOCX with a table."""
    path = tmp_path / "with_tables.docx"
    doc = Document()
    _add_heading(doc, "Data Section", level=1)
    _add_table(doc, [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]])
    doc.save(str(path))
    return path


@pytest.fixture
def docx_mixed_content(tmp_path: Path) -> Path:
    """DOCX with paragraphs, headings, and tables interleaved."""
    path = tmp_path / "mixed.docx"
    doc = Document()
    _add_paragraph(doc, "Intro paragraph.")
    _add_heading(doc, "Section A", level=1)
    _add_paragraph(doc, "Section A text.")
    _add_table(doc, [["Col1", "Col2"], ["V1", "V2"]])
    _add_heading(doc, "Section B", level=2)
    _add_paragraph(doc, "Section B text.")
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx(tmp_path: Path) -> Path:
    """DOCX with no text content at all."""
    path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(path))
    return path


@pytest.fixture
def invalid_docx(tmp_path: Path) -> Path:
    """File with .docx extension but invalid content."""
    path = tmp_path / "invalid.docx"
    path.write_text("This is not a DOCX file.")
    return path
